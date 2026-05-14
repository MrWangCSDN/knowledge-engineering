"""验证 Word 导出（docx_exporter.build_docx）。

测试策略：
  - 用 python-docx 在内存里生成 .docx 字节
  - 用 io.BytesIO + zipfile/python-docx 反序列化看里面内容
  - 关注：文件能开 / 标题在 / 6 段都在 / 没乱码

不打文件系统：build_docx 返回 bytes，直接用 BytesIO 测试。
"""
import io
import zipfile

import pytest

# 待实现的纯函数；现在 import 会失败 → RED
from src.service.qa_engine.docx_exporter import build_docx


# ───────── helpers ─────────


def _make_sections() -> list[dict]:
    """造一份 6 段式答案；跟 SynthesizedAnswer.sections 同形状。"""
    return [
        {
            "type": "overview",
            "title": "业务概述",
            "content": "兽医列表接口的业务概述……",
            "references": [],
        },
        {
            "type": "entry_point",
            "title": "入口方法",
            "content": "VetController.showResourcesVetList()\n  GET /vets",
            "references": [
                {"entity_id": "method//31bc601ec3c2",
                 "display_text": "VetController.showResourcesVetList()",
                 "kind": "method"},
            ],
        },
    ]


def _open_docx(data: bytes):
    """工具：把 bytes 当成 docx 解析回来，方便断言内容。"""
    from docx import Document  # 局部 import 避免文件顶部依赖
    return Document(io.BytesIO(data))


# ───────── RED 1: build_docx 返回非空 bytes ─────────


def test_build_docx_returns_bytes() -> None:
    """最基本：返回值是非空 bytes。"""
    data = build_docx(
        question="兽医列表怎么实现？",
        sections=_make_sections(),
        project_name="Spring PetClinic",
    )
    assert isinstance(data, bytes)
    # docx 文件是 zip；至少有 docx zip magic（前两字节 PK）
    assert data[:2] == b"PK", "应该是 zip-based docx 格式"
    # 大小至少 1KB（一个最简单的 docx 也比这大）
    assert len(data) > 1000


def test_build_docx_is_valid_zip_with_word_document() -> None:
    """打开 zip 应该能找到 word/document.xml（docx 标准结构）。"""
    data = build_docx(
        question="x", sections=_make_sections(), project_name="X",
    )
    with zipfile.ZipFile(io.BytesIO(data)) as z:
        names = z.namelist()
        assert "word/document.xml" in names
        # 主文档 XML 不能为空
        body = z.read("word/document.xml").decode("utf-8")
        assert len(body) > 100


# ───────── RED 2: 问题 + 6 段都进了文档 ─────────


def test_build_docx_includes_question_and_section_titles() -> None:
    """问题文本 + 每个 section 的标题都应该出现在文档里。"""
    sections = [
        {"type": "overview", "title": "业务概述", "content": "概述内容", "references": []},
        {"type": "entry_point", "title": "入口方法", "content": "VetController", "references": []},
        {"type": "call_chain", "title": "调用链路", "content": "A → B → C", "references": []},
        {"type": "db_ops", "title": "数据库操作", "content": "SELECT * FROM vet", "references": []},
        {"type": "rules", "title": "关键约束", "content": "状态校验……", "references": []},
    ]
    data = build_docx(
        question="兽医列表怎么实现？",
        sections=sections,
        project_name="Spring PetClinic",
    )

    doc = _open_docx(data)
    # 把所有段落文本合一起，方便 substring 检查
    full_text = "\n".join(p.text for p in doc.paragraphs)

    # 问题原文必须在
    assert "兽医列表怎么实现？" in full_text
    # 5 个 section 标题都在
    for title in ["业务概述", "入口方法", "调用链路", "数据库操作", "关键约束"]:
        assert title in full_text, f"标题缺失：{title}"
    # 正文内容也在
    assert "VetController" in full_text
    assert "SELECT * FROM vet" in full_text


# ───────── RED 3: references 渲染成项目符号列表 ─────────


def test_build_docx_renders_references_as_bullets() -> None:
    """references 应该用 Word 的 List Bullet 样式渲染（每条一行）。"""
    sections = [
        {
            "type": "entry_point", "title": "入口方法",
            "content": "VetController.showResourcesVetList()",
            "references": [
                {"entity_id": "method//abc", "display_text": "VetController.show()", "kind": "method"},
                {"entity_id": "class//def", "display_text": "VetService", "kind": "class"},
            ],
        },
    ]
    data = build_docx(question="x", sections=sections, project_name="P")
    doc = _open_docx(data)

    # 找出所有 List Bullet 段落
    bullet_paras = [p for p in doc.paragraphs if p.style.name == "List Bullet"]
    # 至少 2 个（我们 2 个 references）
    assert len(bullet_paras) >= 2

    bullet_texts = "\n".join(p.text for p in bullet_paras)
    assert "VetController.show()" in bullet_texts
    assert "VetService" in bullet_texts
    # kind 标签也应该带上
    assert "[method]" in bullet_texts
    assert "[class]" in bullet_texts


# ───────── RED 4: 页眉含项目名 + 生成时间 ─────────


def test_build_docx_header_contains_project_name_and_timestamp() -> None:
    """页眉应该有项目名 + ISO 风格的生成时间。"""
    from datetime import datetime as dt

    fixed_time = dt(2026, 5, 11, 14, 30)
    data = build_docx(
        question="x",
        sections=_make_sections(),
        project_name="存款核心系统",
        generated_at=fixed_time,
    )
    doc = _open_docx(data)

    # docx 的 section.header 默认有一个 paragraphs[0]
    header_para = doc.sections[0].header.paragraphs[0]
    header_text = header_para.text
    assert "存款核心系统" in header_text
    # 时间格式 YYYY-MM-DD
    assert "2026-05-11" in header_text
    assert "14:30" in header_text


# ───────── v1.9: Mermaid → PNG 渲染 ─────────


def test_render_mermaid_to_png_returns_none_when_mmdc_not_installed(monkeypatch):
    """mmdc 不在 PATH 时返回 None（优雅降级）。"""
    from src.service.qa_engine import docx_exporter

    # `shutil.which` 是 docx_exporter 内部用来找 mmdc 的；mock 它返回 None
    monkeypatch.setattr("shutil.which", lambda name: None)
    result = docx_exporter.render_mermaid_to_png("graph LR\n  A --> B")
    assert result is None


def test_render_mermaid_to_png_calls_mmdc_subprocess_and_returns_bytes(monkeypatch, tmp_path):
    """mmdc 可用时：subprocess 调用 + 读取生成的 PNG 文件 → 返回 bytes。"""
    from src.service.qa_engine import docx_exporter

    fake_png_bytes = b"\x89PNG\r\n\x1a\n" + b"fake_png_payload"

    monkeypatch.setattr("shutil.which", lambda name: "/usr/local/bin/mmdc" if name == "mmdc" else None)

    # mock subprocess.run：模拟 mmdc 成功跑完
    # mmdc 调用约定：mmdc -i input.mmd -o output.png
    # 我们的实现应该在 tmp 目录里建 input/output 文件，然后调 subprocess
    import subprocess

    def fake_run(args, **kwargs):
        # 验证调用方式：args[0]='mmdc', '-i'/'-o' 各有一个
        assert "mmdc" in args[0]
        assert "-i" in args
        assert "-o" in args
        # 找到 -o 后面的输出路径，往里写假 PNG
        i = args.index("-o")
        output_path = args[i + 1]
        with open(output_path, "wb") as f:
            f.write(fake_png_bytes)
        # 返回 CompletedProcess(returncode=0)
        return subprocess.CompletedProcess(args=args, returncode=0, stdout=b"", stderr=b"")

    monkeypatch.setattr("subprocess.run", fake_run)

    result = docx_exporter.render_mermaid_to_png("graph LR\n  A --> B")
    assert result == fake_png_bytes


def test_render_mermaid_to_png_returns_none_on_subprocess_failure(monkeypatch):
    """subprocess 报错（exit code != 0）→ 返回 None；不抛错。"""
    from src.service.qa_engine import docx_exporter
    import subprocess

    monkeypatch.setattr("shutil.which", lambda name: "/usr/local/bin/mmdc")

    def fake_run(args, **kwargs):
        return subprocess.CompletedProcess(args=args, returncode=1, stdout=b"", stderr=b"mmdc error")

    monkeypatch.setattr("subprocess.run", fake_run)

    result = docx_exporter.render_mermaid_to_png("invalid mermaid {{{")
    assert result is None


def test_build_docx_calls_render_mermaid_to_png_for_call_chain_section(monkeypatch):
    """call_chain 段含 ```mermaid``` 时 build_docx 应该调 render_mermaid_to_png。"""
    from src.service.qa_engine import docx_exporter

    # 记录 render 被以什么 code 调用
    captured: list[str] = []

    def fake_render(code: str):
        captured.append(code)
        # 返回 None 让 build_docx 走 fallback 代码块路径，不踩 PNG 嵌入坑
        return None

    monkeypatch.setattr(docx_exporter, "render_mermaid_to_png", fake_render)

    sections = [
        {
            "type": "call_chain",
            "title": "调用链",
            "content": "前文\n```mermaid\ngraph LR\n  A --> B\n```\n后文",
            "references": [],
        },
    ]
    build_docx(
        question="x",
        sections=sections,
        project_name="P",
    )

    # render_mermaid_to_png 应该被调一次，参数是 fence 里的 mermaid 源码
    assert len(captured) == 1
    assert "graph LR" in captured[0]
    assert "A --> B" in captured[0]


def test_build_docx_falls_back_to_code_block_when_png_render_fails(monkeypatch):
    """render_mermaid_to_png 返回 None 时，mermaid code 应该作为文本保留在文档里（fallback）。"""
    from src.service.qa_engine import docx_exporter

    monkeypatch.setattr(docx_exporter, "render_mermaid_to_png", lambda code: None)

    sections = [
        {
            "type": "call_chain",
            "title": "调用链",
            "content": "```mermaid\ngraph LR\n  Owner --> Pet\n```",
            "references": [],
        },
    ]
    data = build_docx(
        question="x",
        sections=sections,
        project_name="P",
    )
    doc = _open_docx(data)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    # mermaid 源码作为代码段保留
    assert "graph LR" in full_text
    assert "Owner --> Pet" in full_text


def test_build_docx_skips_mermaid_render_for_non_call_chain_sections(monkeypatch):
    """non-call_chain 段即便含 ```mermaid 也不调 render（只 call_chain 段触发图嵌入）。"""
    from src.service.qa_engine import docx_exporter

    captured: list[str] = []
    monkeypatch.setattr(docx_exporter, "render_mermaid_to_png",
                        lambda code: captured.append(code) or None)

    sections = [
        # overview 段含 mermaid fence —— 不应该触发 render
        {
            "type": "overview",
            "title": "业务概述",
            "content": "概述里偶然有 ```mermaid\ngraph A\n``` 块",
            "references": [],
        },
    ]
    build_docx(question="x", sections=sections, project_name="P")
    assert captured == []  # 一次都没调


# ───────── v1.2: chit-chat fallback ─────────


def test_docx_exporter_handles_chit_chat_section_type():
    """v1.2: chit-chat session 也可以导 docx，title/icon 有兜底。"""
    from src.service.qa_engine.docx_exporter import _SECTION_TITLES, _SECTION_EMOJIS
    assert "chit-chat" in _SECTION_TITLES
    assert "chit-chat" in _SECTION_EMOJIS
    assert _SECTION_TITLES["chit-chat"] == "对话回复"
    assert _SECTION_EMOJIS["chit-chat"] == "💬"

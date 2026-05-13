"""验证 Word 模板机制（v1.9）。

用户上传一份带占位符的 .docx 模板（自己加 logo / 页眉 / 品牌色），
build_docx 读取它替换占位符填进 sections，输出一份个性化 docx。

模板占位符约定：
  {{QUESTION}}        ← 用户问题
  {{PROJECT_NAME}}    ← 工程名
  {{GENERATED_AT}}    ← 生成时间（ISO 8601 字符串）
  {{SECTIONS}}        ← 一整块"6 段答案"会替换这一占位段；
                       该段所在的 paragraph 会被删除，原位置插入 6 段答案的 paragraphs

环境变量 / 参数：
  build_docx_from_template(template_path, question, sections, project_name, generated_at)
"""
import io
import zipfile

import pytest
from docx import Document
from docx.shared import Pt

# 待实现的新工厂
from src.service.qa_engine.docx_exporter import build_docx_from_template


@pytest.fixture
def simple_template_path(tmp_path):
    """造一份最小可用模板：包含 4 个占位符的 .docx。

    `tmp_path` 是 pytest 自带 fixture，每个测试一个独立目录，自动清理。
    """
    doc = Document()
    # H1 用 Word 内置 Title 风格让人能区分（模板里通常有 logo / 公司名等品牌元素）
    doc.add_heading("XX 银行 · 代码知识问答记录", level=0)
    # 普通段落 + 占位符
    doc.add_paragraph("问题：{{QUESTION}}")
    doc.add_paragraph("工程：{{PROJECT_NAME}}    生成时间：{{GENERATED_AT}}")
    # SECTIONS 占位整段 —— 替换为 6 段答案
    doc.add_paragraph("{{SECTIONS}}")
    # 模板下方有"机密 仅供参考"等品牌文案，应被保留
    doc.add_paragraph("⚠️ 本文档为内部资料，仅限授权人员阅读。")

    out = tmp_path / "template.docx"
    doc.save(out)
    return str(out)


def _open_docx(data: bytes) -> Document:
    return Document(io.BytesIO(data))


def test_build_docx_from_template_replaces_simple_placeholders(simple_template_path):
    """{{QUESTION}} / {{PROJECT_NAME}} / {{GENERATED_AT}} 应该被实际值替换。"""
    from datetime import datetime
    data = build_docx_from_template(
        template_path=simple_template_path,
        question="兽医列表怎么实现？",
        sections=[{"type": "overview", "title": "业务概述", "content": "答案", "references": []}],
        project_name="Spring PetClinic",
        generated_at=datetime(2026, 5, 11, 14, 30),
    )
    doc = _open_docx(data)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    # 真值应替换占位符
    assert "兽医列表怎么实现？" in full_text
    assert "Spring PetClinic" in full_text
    assert "2026-05-11" in full_text
    # 占位符不应残留
    assert "{{QUESTION}}" not in full_text
    assert "{{PROJECT_NAME}}" not in full_text
    assert "{{GENERATED_AT}}" not in full_text


def test_build_docx_from_template_inserts_sections_at_placeholder(simple_template_path):
    """{{SECTIONS}} 占位段应该被 6 段答案取代（标题 + 内容）。"""
    sections = [
        {"type": "overview", "title": "业务概述", "content": "概述内容", "references": []},
        {"type": "entry_point", "title": "入口方法", "content": "VetController", "references": []},
        {"type": "rules", "title": "关键约束", "content": "约束规则", "references": []},
    ]
    data = build_docx_from_template(
        template_path=simple_template_path,
        question="x",
        sections=sections,
        project_name="P",
    )
    doc = _open_docx(data)
    full_text = "\n".join(p.text for p in doc.paragraphs)

    # 3 段标题 + 3 段内容都应该出现
    for title in ["业务概述", "入口方法", "关键约束"]:
        assert title in full_text
    for content in ["概述内容", "VetController", "约束规则"]:
        assert content in full_text
    # SECTIONS 占位符不应残留
    assert "{{SECTIONS}}" not in full_text


def test_build_docx_from_template_preserves_brand_elements(simple_template_path):
    """模板里 logo 区 / 品牌文案应该保留下来（不被覆盖）。"""
    data = build_docx_from_template(
        template_path=simple_template_path,
        question="x",
        sections=[{"type": "overview", "title": "x", "content": "y", "references": []}],
        project_name="P",
    )
    doc = _open_docx(data)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    # 模板第 1 段：H0 标题（XX 银行）
    assert "XX 银行" in full_text
    # 模板尾段：保密文案
    assert "机密" in full_text or "内部资料" in full_text


def test_build_docx_from_template_nonexistent_path_raises(tmp_path):
    """模板文件不存在 → 抛 FileNotFoundError（早暴露配置错误）。"""
    missing = tmp_path / "no-such-template.docx"
    with pytest.raises(FileNotFoundError):
        build_docx_from_template(
            template_path=str(missing),
            question="x", sections=[], project_name="P",
        )

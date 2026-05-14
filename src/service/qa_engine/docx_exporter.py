"""把 QA 答案导出成企业级 Word 模板。

依赖：python-docx 1.x

设计原则（[[首页设计]] §16 v1.5）：
  - **纯函数**：build_docx(question, sections, project_name) → bytes
    没有副作用、不写文件，便于测试 + HTTP endpoint 直接 stream
  - **模板化**：标题 / 页眉 / 页脚 / 段落样式都从一个地方控制
  - **Mermaid 暂不嵌入**：v1.5 把 ```mermaid 块作为代码块原样塞进去；
    v1.6 再接 mermaid-cli 转 PNG 嵌入（这是用户最后期望的"企业级"目标）
"""
from __future__ import annotations

# io：标准库，用 BytesIO 当内存"假文件"，让 Document.save 写到 bytes 里
import io
import os
# re：从 call_chain 段抽 ```mermaid fence
import re
# shutil.which 判断 mmdc 在 PATH 里没；subprocess 跑外部命令
import shutil
import subprocess
# tempfile：mmdc 需要文件 input + 文件 output，用 tempdir 隔离每次调用
import tempfile
from copy import deepcopy
from datetime import datetime
from typing import Any, Optional

# python-docx 主入口；Document 是 Word 文档对象
from docx import Document
# Pt = 字号单位（point，1 pt = 1/72 inch）；Cm = 厘米；用来设页边距
from docx.shared import Pt, Cm, Inches
# 段落对齐枚举：LEFT / CENTER / RIGHT / JUSTIFY
from docx.enum.text import WD_ALIGN_PARAGRAPH


# 段类型 → 中文标题（覆盖 LLM 没给 title 的兜底）
_SECTION_TITLES: dict[str, str] = {
    "overview": "业务概述",
    "entry_point": "入口方法",
    "call_chain": "调用链路",
    "db_ops": "数据库操作",
    "rules": "关键约束与业务规则",
    "sources": "引用来源",
    "chit-chat": "对话回复",  # v1.2: 防 KeyError 兜底（chit-chat 一般不导 docx）
}

# 段类型 → emoji（仅在 Word 的小标题里点缀；不依赖 emoji 也能看）
_SECTION_EMOJIS: dict[str, str] = {
    "overview": "📋",
    "entry_point": "🚪",
    "call_chain": "🔀",
    "db_ops": "💾",
    "rules": "⚠️",
    "sources": "🔗",
    "chit-chat": "💬",  # v1.2
}


def build_docx(
    *,
    question: str,
    sections: list[dict[str, Any]],
    project_name: str,
    generated_at: datetime | None = None,
) -> bytes:
    """把 QA 答案构造成一份 .docx，返回二进制内容。

    :param question:     用户的原始问题（出现在文档标题下方）
    :param sections:     6 段式 sections 列表（来自 SynthesizedAnswer.sections）
    :param project_name: 工程名，写到页眉
    :param generated_at: 生成时间；默认 now()
    :return: docx 二进制 bytes，可直接做 HTTP response body 或写文件
    """
    # 1) 构造 Document：默认有 1 个 section（不是答案的 section，是 Word 的页面 section）
    doc = Document()

    # 2) 设页边距（A4 视觉舒适）
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    # 3) 页眉：项目名 + 生成时间（右对齐）
    _setup_header(doc, project_name=project_name, generated_at=generated_at or datetime.now())

    # 4) 页脚：版权占位（v1.6 让用户配置）
    _setup_footer(doc)

    # 5) 主标题：用户问题
    title = doc.add_heading("代码知识问答记录", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 6) 问题段（次标题样式）
    q_para = doc.add_paragraph()
    q_para.add_run("问题：").bold = True
    # 多个 run 才能在一段里分粗细字
    q_para.add_run(question)

    # 7) 一段空行分隔
    doc.add_paragraph()

    # 8) 依次写每个 section
    for s in sections:
        _write_section(doc, s)

    # 9) 序列化到内存
    # BytesIO 是文件对象语义的内存容器；Document.save 接受路径 OR 文件对象
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── 内部 helpers ─────────────────────────────────────────────────────────


def build_docx_from_template(
    *,
    template_path: str,
    question: str,
    sections: list[dict[str, Any]],
    project_name: str,
    generated_at: datetime | None = None,
) -> bytes:
    """从用户自带的 .docx 模板生成 Word 文档（v1.9）。

    模板里允许使用占位符（必须在普通段落里，不能在表格 / 页眉里）：
      {{QUESTION}}        ← 替换成用户问题
      {{PROJECT_NAME}}    ← 替换成工程名
      {{GENERATED_AT}}    ← 替换成生成时间字符串
      {{SECTIONS}}        ← 替换成一整块 6 段答案 paragraphs

    实现细节：
      - {{X}} 简单占位符 → 段落内字符串 replace
      - {{SECTIONS}} → 找到该段，在原位置插入多个 paragraphs（标题 + 内容），删除占位段

    :param template_path: .docx 模板绝对路径；不存在抛 FileNotFoundError
    :return: 渲染后的 docx 二进制
    """
    # 路径存在性早期校验，比 python-docx 内部抛个晦涩异常友好
    if not os.path.isfile(template_path):
        raise FileNotFoundError(f"docx 模板不存在: {template_path}")

    # 加载模板；这一步会把模板里的 styles / 页眉 / 页脚 / logo 都保留
    doc = Document(template_path)

    gen_at = generated_at or datetime.now()
    generated_at_str = gen_at.strftime("%Y-%m-%d %H:%M")

    # ─── 第一步：简单占位符替换（保留段落格式）───
    # 简单占位符 = {{X}} → 字符串值
    simple_replacements = {
        "{{QUESTION}}": question,
        "{{PROJECT_NAME}}": project_name,
        "{{GENERATED_AT}}": generated_at_str,
    }
    # 遍历所有 paragraph；每段内的所有 run 拼起来做 replace
    # 注意：python-docx 的 paragraph 由多个 run（行内片段）组成，
    # 跨 run 的占位符可能会被切断；这里用简单实现处理"整段单 run"的常见情况
    for paragraph in list(doc.paragraphs):
        full = paragraph.text
        modified = full
        for placeholder, value in simple_replacements.items():
            if placeholder in modified:
                modified = modified.replace(placeholder, value)
        if modified != full:
            # 清空原 run 后用单个 run 重写
            # 简化做法：删所有 run，再 add 一个 run；丢失了细粒度样式但保留段落级样式
            for run in list(paragraph.runs):
                run.text = ""
            # 把替换后的整段写在第一个 run 里
            if paragraph.runs:
                paragraph.runs[0].text = modified
            else:
                paragraph.add_run(modified)

    # ─── 第二步：{{SECTIONS}} 替换为多段 paragraphs ───
    sections_placeholder = "{{SECTIONS}}"
    # 找占位符段；用 list() 复制因为我们会修改 doc.paragraphs
    for paragraph in list(doc.paragraphs):
        if sections_placeholder in paragraph.text:
            # 在该段位置插入 sections，最后删除占位段
            _insert_sections_before(paragraph, sections)
            # 删除占位段（用 XML 操作；python-docx 没有 delete API）
            _delete_paragraph(paragraph)
            break  # 只替换第一个匹配，避免重复插入

    # ─── 序列化 ───
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _insert_sections_before(anchor_paragraph, sections: list[dict[str, Any]]) -> None:
    """在 anchor_paragraph 前插入 sections 渲染的 paragraphs。

    python-docx 没有直接的"在某段前插入新段"API，所以用底层 XML 操作。
    """
    # 拿到底层 XML 元素 ref
    anchor_elem = anchor_paragraph._element
    # 父元素（body）
    parent = anchor_elem.getparent()
    # 当前 anchor 在父中的 index
    anchor_idx = list(parent).index(anchor_elem)

    # 为了用 python-docx 高阶 API 添加段落（而不是手写 XML），
    # 先把段落 append 到 doc 末尾，再把它们的底层 XML 移到 anchor 前
    # 这是 python-docx 社区最常用的"在中间插入"技巧
    doc = anchor_paragraph.part.document

    # 收集新建段落的 XML 元素
    new_elements = []
    for section in sections:
        section_type = section.get("type", "")
        title = section.get("title") or _SECTION_TITLES.get(section_type, section_type)
        emoji = _SECTION_EMOJIS.get(section_type, "")

        # 小标题段
        title_para = doc.add_heading(f"{emoji} {title}".strip(), level=2)
        new_elements.append(title_para._element)

        # 正文段（按 \n 拆多段）
        content = section.get("content", "") or ""
        for line in content.split("\n"):
            content_para = doc.add_paragraph(line)
            new_elements.append(content_para._element)

        # 引用列表
        refs = section.get("references") or []
        if refs:
            italic_para = doc.add_paragraph()
            italic_para.add_run("引用：").italic = True
            new_elements.append(italic_para._element)
            for r in refs:
                display_text = r.get("display_text") or r.get("entity_id", "?")
                kind = r.get("kind", "")
                bullet_para = doc.add_paragraph(f"{display_text}  [{kind}]", style="List Bullet")
                new_elements.append(bullet_para._element)

    # 把这些刚 append 到末尾的元素，整体搬到 anchor 之前
    # `parent.remove + insert` 是 lxml 的标准移动方式
    for offset, elem in enumerate(new_elements):
        parent.remove(elem)
        parent.insert(anchor_idx + offset, elem)


def _delete_paragraph(paragraph) -> None:
    """从父元素删除一个 paragraph（python-docx 没有官方 API）。"""
    elem = paragraph._element
    elem.getparent().remove(elem)


def _setup_header(doc, *, project_name: str, generated_at: datetime) -> None:
    """页眉：左侧项目名 + 右侧时间戳。

    Word 页眉一般是 "section.header.paragraphs[0]"。
    制表符切左右对齐是企业模板的常用技巧。
    """
    # 取第一个 section 的 header；docx 默认是单 section 文档
    header = doc.sections[0].header
    p = header.paragraphs[0]
    # 左 + tab + 右；后面会用 right_tab_stop 让 tab 后的文本右对齐
    p.text = f"{project_name}\t\t代码知识工程 · {generated_at:%Y-%m-%d %H:%M}"
    # 改字号 / 字色（页眉应该比正文小一点）
    for run in p.runs:
        run.font.size = Pt(9)


def _setup_footer(doc) -> None:
    """页脚：占位『生成自代码知识工程 v1.5』。"""
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.text = "本文档由代码知识工程自动生成 · 仅供参考"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(8)


def _write_section(doc, section: dict[str, Any]) -> None:
    """写一个段（标题 + 内容 + 引用列表）。

    v1.9: call_chain 段含 ```mermaid``` 时，能转 PNG 就嵌图，否则当代码块。
    """
    section_type = section.get("type", "")
    title = section.get("title") or _SECTION_TITLES.get(section_type, section_type)
    emoji = _SECTION_EMOJIS.get(section_type, "")

    # H2 小标题（用 Word 内置 Heading 2 样式）
    doc.add_heading(f"{emoji} {title}".strip(), level=2)

    # 正文：v1.9 起对 call_chain 做 mermaid fence 检测
    content = section.get("content", "") or ""
    if content:
        _write_section_content(doc, content, section_type)

    # 引用列表（如果有）
    refs = section.get("references") or []
    if refs:
        # 加个小标题
        doc.add_paragraph().add_run("引用：").italic = True
        for r in refs:
            display_text = r.get("display_text") or r.get("entity_id", "?")
            kind = r.get("kind", "")
            # List Bullet 是 Word 内置项目符号样式
            doc.add_paragraph(f"{display_text}  [{kind}]", style="List Bullet")


# 匹配 ```mermaid ... ``` fence；(?s) 让 . 跨行匹配；非贪婪
_MERMAID_FENCE_RE = re.compile(r"```mermaid\s*\n(.*?)\n```", re.DOTALL)


def _write_section_content(doc, content: str, section_type: str) -> None:
    """把 section.content 写到文档里；call_chain 段尝试嵌入 Mermaid PNG。"""
    # 非 call_chain 段直接按行写
    if section_type != "call_chain":
        for line in content.split("\n"):
            doc.add_paragraph(line)
        return

    # call_chain 段：扫 ```mermaid fence；fence 之外的文字按行写
    last_end = 0
    for match in _MERMAID_FENCE_RE.finditer(content):
        # fence 之前的文字
        before = content[last_end: match.start()]
        for line in before.split("\n"):
            if line.strip():  # 跳过空行避免段落过多
                doc.add_paragraph(line)

        # fence 内的 mermaid code → 尝试转 PNG
        mermaid_code = match.group(1)
        png_bytes = render_mermaid_to_png(mermaid_code)
        if png_bytes is not None:
            # 嵌图：python-docx 接受 BytesIO；宽度限制 6 英寸防止超页
            try:
                doc.add_picture(io.BytesIO(png_bytes), width=Inches(6))
            except Exception:
                # 万一 png 格式 / 注入有问题，回退代码块
                _add_code_block(doc, mermaid_code)
        else:
            # 转 PNG 失败 → fallback 代码块（v1.5 行为）
            _add_code_block(doc, mermaid_code)

        last_end = match.end()

    # fence 之后的尾巴
    tail = content[last_end:]
    if tail.strip():
        for line in tail.split("\n"):
            if line.strip():
                doc.add_paragraph(line)


def _add_code_block(doc, code: str) -> None:
    """把代码块当 monospace 段落塞进 doc（fallback 走这个）。"""
    para = doc.add_paragraph()
    run = para.add_run(code)
    # Courier New / Consolas 是常见 monospace
    run.font.name = "Courier New"
    run.font.size = Pt(9)


# ─── v1.9 Mermaid PNG 渲染 ──────────────────────────────────────────────


def render_mermaid_to_png(code: str) -> Optional[bytes]:
    """用 mermaid-cli（mmdc）把 mermaid 源码转 PNG bytes。

    优雅降级：mmdc 不在 PATH 里 → 返回 None；调用方需要做兜底（写代码块等）。

    :param code: mermaid 源码（不含 ```mermaid fence 标记）
    :return: PNG bytes 或 None
    """
    # 找 mmdc 可执行文件
    mmdc_path = shutil.which("mmdc")
    if not mmdc_path:
        return None

    # mmdc 要求文件输入 + 文件输出；用 tempdir 隔离
    try:
        with tempfile.TemporaryDirectory(prefix="ke-mermaid-") as tmpdir:
            in_path = os.path.join(tmpdir, "in.mmd")
            out_path = os.path.join(tmpdir, "out.png")
            with open(in_path, "w", encoding="utf-8") as f:
                f.write(code)

            # mmdc -i input.mmd -o output.png
            # 加 -b transparent 防止默认黑底；-w 1200 限定宽度让图清晰
            result = subprocess.run(
                [mmdc_path, "-i", in_path, "-o", out_path, "-b", "transparent", "-w", "1200"],
                capture_output=True,
                # mmdc 偶尔会卡 puppeteer，给 30s 上限
                timeout=30,
            )
            if result.returncode != 0:
                return None
            if not os.path.isfile(out_path):
                return None
            with open(out_path, "rb") as f:
                return f.read()
    except Exception:
        # 任何异常（subprocess 启动失败 / 超时 / IO 错）都兜底 None
        return None
